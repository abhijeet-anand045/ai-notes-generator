import streamlit as st
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
import os
import docx
from io import BytesIO  


llm = ChatGroq(temperature=0.5, groq_api_key="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", model_name="llama3-8b-8192")


# Placeholder functions to simulate LLM invocation (replace with actual LLM code)
def notes_gen(subject_name, syllabus, level):
    prompt_template = """
    You are an expert in creating educational content. Based on the following inputs, generate Very detailed notes for student of {level} to get very in depth knowledge.

    Subject: {subject_name}
    Syllabus: {syllabus}

    Please cover all topics as possible sub-topics given in the {syllabus}
    And please make it a proper structured    
    """

    prompt = prompt_template.format(
        subject_name=subject_name,
        syllabus=syllabus,
        level=level
    )

    chain = (
        ChatPromptTemplate.from_template(prompt)
        | llm
        | StrOutputParser()
    )

    return chain.invoke({})

if "notes_gen" not in st.session_state:
    st.session_state.notes_gen = ""
# Streamlit app code
st.title("Notes Generation")


# Sidebar inputs
st.sidebar.header("Input Details")

# Subject Name
subject_name = st.sidebar.text_input("Subject Name")
if subject_name:
    st.sidebar.markdown("✅ Subject Name entered")

# Syllabus
syllabus = st.sidebar.text_area("Syllabus")
if syllabus:
    st.sidebar.markdown("✅ Syllabus entered")

#Level
level=st.radio("Select Your Level for the notes:",("Class 10 or below","Class 11 or 12","Undergradute Enterance","Undergradute","Postgradute Enterance","Postgraduate","Goverment competitive exams"))

# Button to generate Notes
if st.sidebar.button("Generate Notes"):
    if subject_name and syllabus and level:
        st.session_state.notes_gen = notes_gen(subject_name, syllabus, level)
        st.header("Generated Notes")
        st.text_area("Generated Notes",value=st.session_state.notes_gen, height=400)
    else:
        st.sidebar.markdown("❗ Please fill in all required fields to generate Notes")

# Button to download questions as a DOCX file
if st.sidebar.button("Generate All Questions as DOCX"):
    if st.session_state.notes_gen:
        doc = docx.Document()
        doc.add_heading("Generated Notes", level=1)
        doc.add_paragraph(st.session_state.notes_gen)
        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Offer the file for download
        st.download_button(
            label="Download Generated Questions as DOCX",
            data=doc_io,
            file_name="generated_notes.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.sidebar.markdown("❗ Generate some Notes first before downloading")